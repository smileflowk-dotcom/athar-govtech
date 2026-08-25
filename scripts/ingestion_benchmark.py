from __future__ import annotations

import csv
import json
import time
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
FIXTURES = ROOT / "tmp" / "ingestion-benchmark"
OUT = ROOT / "tmp" / "ingestion-results"
FIXTURES.mkdir(parents=True, exist_ok=True)
OUT.mkdir(parents=True, exist_ok=True)

EXPECTED_PHRASES = [
    "Contrôle de la commande publique",
    "Chaque alerte mène à sa preuve",
    "publication des offres",
    "validation humaine",
]


def make_fixtures() -> list[Path]:
    from PIL import Image, ImageDraw, ImageFont
    from docx import Document
    from openpyxl import Workbook
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    body = [
        "Contrôle de la commande publique",
        "Chaque alerte mène à sa preuve.",
        "Le dossier contient les pièces de publication des offres, d'évaluation et d'attribution.",
        "Les contrôles sont explicables et soumis à validation humaine.",
        "Point à vérifier : délai entre publication et ouverture des plis.",
    ]

    native_pdf = FIXTURES / "native-fr.pdf"
    c = canvas.Canvas(str(native_pdf), pagesize=A4)
    y = 790
    for idx, line in enumerate(body):
        c.setFont("Helvetica-Bold" if idx == 0 else "Helvetica", 15 if idx == 0 else 11)
        c.drawString(60, y, line)
        y -= 42
    c.showPage()
    c.save()

    scan_png = FIXTURES / "scan-fr.png"
    img = Image.new("RGB", (1654, 2339), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 38)
        title_font = ImageFont.truetype("DejaVuSans-Bold.ttf", 48)
    except Exception:
        font = ImageFont.load_default()
        title_font = font
    y = 150
    for idx, line in enumerate(body):
        draw.text((120, y), line, fill="black", font=title_font if idx == 0 else font)
        y += 120
    img.save(scan_png, quality=85)
    scan_pdf = FIXTURES / "scan-fr.pdf"
    img.save(scan_pdf, "PDF", resolution=150.0)

    docx_path = FIXTURES / "dossier.docx"
    doc = Document()
    doc.add_heading(body[0], level=1)
    for line in body[1:]:
        doc.add_paragraph(line)
    doc.save(docx_path)

    xlsx_path = FIXTURES / "offres.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Offres"
    ws.append(["Soumissionnaire", "Montant MAD", "Classement"])
    ws.append(["Alpha", 1200000, 2])
    ws.append(["Beta", 1150000, 1])
    ws.append(["Gamma", 1300000, 3])
    wb.save(xlsx_path)

    return [native_pdf, scan_pdf, docx_path, xlsx_path]


def flatten_text(obj: Any) -> str:
    chunks: list[str] = []
    if isinstance(obj, dict):
        for key, value in obj.items():
            if key in {"text", "orig", "content"} and isinstance(value, str):
                chunks.append(value)
            else:
                chunks.append(flatten_text(value))
    elif isinstance(obj, list):
        for value in obj:
            chunks.append(flatten_text(value))
    return "\n".join(x for x in chunks if x)


def count_key(obj: Any, target: str) -> int:
    if isinstance(obj, dict):
        return sum((1 if k == target else 0) + count_key(v, target) for k, v in obj.items())
    if isinstance(obj, list):
        return sum(count_key(v, target) for v in obj)
    return 0


def accuracy(text: str) -> float:
    lowered = text.lower()
    hits = sum(1 for phrase in EXPECTED_PHRASES if phrase.lower() in lowered)
    return round(hits / len(EXPECTED_PHRASES), 3)


def run_unstructured(path: Path) -> dict[str, Any]:
    from unstructured.partition.auto import partition

    started = time.perf_counter()
    elements = partition(filename=str(path))
    elapsed = time.perf_counter() - started
    data = [e.to_dict() for e in elements]
    text = "\n".join((e.get("text") or "") for e in data)
    pages = sorted({(e.get("metadata") or {}).get("page_number") for e in data if (e.get("metadata") or {}).get("page_number")})
    coords = sum(1 for e in data if (e.get("metadata") or {}).get("coordinates"))
    tables = sum(1 for e in data if e.get("type") == "Table")
    return {
        "engine": "unstructured-oss",
        "file": path.name,
        "success": True,
        "seconds": round(elapsed, 3),
        "elements": len(data),
        "pages": len(pages),
        "coordinates": coords,
        "tables": tables,
        "text_chars": len(text),
        "expected_phrase_accuracy": accuracy(text),
        "sample": text[:1200],
    }


def run_docling(path: Path) -> dict[str, Any]:
    from docling.document_converter import DocumentConverter

    started = time.perf_counter()
    result = DocumentConverter().convert(str(path))
    elapsed = time.perf_counter() - started
    exported = result.document.export_to_dict()
    text = result.document.export_to_markdown()
    prov = count_key(exported, "prov")
    tables = len(exported.get("tables", [])) if isinstance(exported, dict) else 0
    pages = len(exported.get("pages", {})) if isinstance(exported, dict) and isinstance(exported.get("pages"), dict) else 0
    return {
        "engine": "docling-oss",
        "file": path.name,
        "success": True,
        "seconds": round(elapsed, 3),
        "elements": count_key(exported, "text"),
        "pages": pages,
        "coordinates": prov,
        "tables": tables,
        "text_chars": len(text),
        "expected_phrase_accuracy": accuracy(text),
        "sample": text[:1200],
    }


def main() -> None:
    files = make_fixtures()
    results: list[dict[str, Any]] = []
    for path in files:
        for runner in (run_unstructured, run_docling):
            try:
                row = runner(path)
            except Exception as exc:
                row = {
                    "engine": "unstructured-oss" if runner is run_unstructured else "docling-oss",
                    "file": path.name,
                    "success": False,
                    "error": f"{type(exc).__name__}: {exc}",
                }
            print(json.dumps(row, ensure_ascii=False))
            results.append(row)

    (OUT / "results.json").write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    fields = ["engine", "file", "success", "seconds", "elements", "pages", "coordinates", "tables", "text_chars", "expected_phrase_accuracy", "error"]
    with (OUT / "results.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(results)


if __name__ == "__main__":
    main()
